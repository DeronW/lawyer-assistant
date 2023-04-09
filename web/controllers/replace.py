from docx import Document
from openpyxl import load_workbook
import operator
from IPython import embed
from io import BytesIO


def gen_from_template(docx_stream, xlsx_stream):
    doc = Document(docx_stream)
    vars = get_vars(xlsx_stream)
    print(vars)

    replace_paragraphs(doc.paragraphs, vars)

    # 处理表格中的文字内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_paragraphs(cell.paragraphs, vars)
    bs = BytesIO()
    doc.save(bs)
    bs.seek(0)  # 返回 0 ，否则 read 方法会返回空
    return bs


def replace_paragraphs(paragraphs, vars):
    for k, v in vars.items():
        if v is None:
            continue
        for paragraph in paragraphs:
            # paragraphs 中的 run 之间是分隔的，并且样式绑定在 run object 上
            if operator.contains(paragraph.text, k):
                # 这里需要注意，不是字符串直接匹配，需要逐个 run object 拼接对比
                start = 0
                t = ""
                for i in range(len(paragraph.runs)):
                    run = paragraph.runs[i]
                    t += run.text
                    if operator.contains(t, k):
                        paragraph.runs[start].text = t.replace(k, v)
                        # 替换后，把多余的 runs 去掉
                        for j in range(start + 1, i + 1):
                            paragraph.runs[j].text = ""
                    elif endswith_partial_target(t, k):
                        # 匹配过程进行中，继续匹配
                        pass
                    else:
                        start = i
                        t = ""


def endswith_partial_target(text, target) -> bool:
    """结尾是否部分包含目标字符串"""
    for i in range(len(target)):
        if text.endswith(target[0 : i + 1]):
            return True
    return False


def get_vars(stream):
    wb = load_workbook(stream)
    sheet = wb[wb.get_sheet_names()[0]]
    vars = {}
    for r in range(1, sheet.max_row + 1):
        k = sheet.cell(r, 1).value
        v = sheet.cell(r, 2).value
        if k.startswith("【"):
            vars[k] = v
    return vars
